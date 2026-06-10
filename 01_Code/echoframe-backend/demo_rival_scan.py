"""
Rival Scan — one-page SAMPLE report generator.
Reuses EchoFrame house-style primitives from engine.py (colors, navy section
bars, KPI tiles, cell helpers, footer) so the sample matches the locked
Clarity Report look. Fictional Columbus GA pizza restaurant + 3 local rivals.
Marked REDACTED SAMPLE. Single page, no real data.
"""
import io
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import engine as E
from engine import (
    NAVY, AMBER, WHITE, RED, GREEN, GRAY,
    NAVY_HEX, AMBER_HEX, LTBLUE_HEX,
    _set_doc_defaults, _hide_table_borders, _get_or_add_tblPr,
    _set_cell_bg, _set_cell_padding, _set_cell_valign, _set_cell_width_dxa,
    _dark_section_bar, _add_bottom_border, _build_footer,
)

SLATE_HEX = "F1F5F9"

# ── Fictional data (sample only) ───────────────────────────────────────────────
META = {
    "Business Name": "Tony's Brick Oven Pizza",
    "Tagline": "Rival Scan  ·  Local Competitive Monitoring  ·  June 2026",
    "Location": "Broadway, Columbus GA",
    "Month": "June 2026",
}

KPIS = [
    ("Competitors Tracked", "3",          "Within 4 mi of Broadway", NAVY),
    ("Changes This Month",  "7",          "▲ up from 4 in May",  RED),
    ("Your Price Position", "Mid-Market", "2nd-lowest of 4 tracked",  NAVY),
]

# Competitor | Key Price (Large 1-Topping) | Rating / Reviews | Promotion | Change
COMP_ROWS = [
    ("Tony's Brick Oven  (You)", "$15.99", "4.6★  (412)",   "None active",                     "—", GRAY),
    ("Marco's Pizza – Macon Rd", "$13.99", "4.1★  (530)", "Online large 1-top $9.99 (PIZZA10)", "▼ –$2.00", RED),
    ("Mellow Mushroom",         "$18.50", "4.3★  (1,240)", "2 for $24 large cheese (ends 6/15)", "▲ +$1.00", RED),
    ("Fountain City Pizza Co.", "$16.50", "4.7★  (286)",   "Free garlic knots over $25",        "—", GRAY),
]

# (marker_hex, bold_lead, rest)
ALERTS = [
    (RED,   "Marco's slashed online pricing.",
     " Large 1-topping dropped to $9.99 with code PIZZA10 — now undercuts you by $6.00 on delivery apps."),
    (AMBER, "Mellow Mushroom raised base + launched a bundle.",
     " Large specialty up $1.00 to $18.50, paired with a “2-for-$24 cheese” promo running through 6/15."),
    (GREEN, "Fountain City's new threshold plays to you.",
     " Their “free garlic knots over $25” kicks in above your ~$25.50 average ticket — easy to out-message."),
    (NAVY,  "Reviews shifting.",
     " 14 new Google reviews across rivals this week; Mellow Mushroom's rating slipped 4.4 → 4.3."),
]

ONE_THING_HEAD = "Counter Marco's $9.99 online push — don't match it."
ONE_THING_BODY = (
    "Marco's loss-leader targets price-sensitive online orders, not your dine-in base. Matching $9.99 "
    "torches margin; ignoring it cedes weekday delivery. Launch a “Tuesday Online-Only Large 1-Topping — $11.99” "
    "fenced to your own app and web ordering, capped at one per order."
)
ONE_THING_STEPS = [
    "Add the $11.99 Tuesday code to your online menu and pin it to the app home screen.",
    "Push it once to your SMS/email list framed as “skip the apps, order direct.”",
    "Leave dine-in and your $15.99 everyday price untouched — this is a fence, not a price cut.",
]
ONE_THING_IMPACT = (
    "Estimated impact: defends ~$1,150/month in at-risk online revenue "
    "(≈35 weekday online orders/wk × $15.99, the cohort most likely to chase Marco's code)."
)
ONE_THING_RISK = "Cannibalization if un-fenced. Keep it online-only and one-per-order so walk-in traffic doesn't trade down."


# ── Custom header (Rival Scan variant of the navy dashboard band) ──────────────
def _build_rival_header(doc, meta):
    biz = meta["Business Name"]
    tbl = doc.add_table(rows=1, cols=2)
    tbl.allow_autofit = False
    _hide_table_borders(tbl)
    tblPr = _get_or_add_tblPr(tbl._tbl)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(int(7.4 * 1440)))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)

    lc, rc = tbl.rows[0].cells
    _set_cell_width_dxa(lc, int(7.4 * 1440 * 0.62))
    _set_cell_width_dxa(rc, int(7.4 * 1440 * 0.38))
    _set_cell_bg(lc, NAVY_HEX)
    _set_cell_bg(rc, NAVY_HEX)
    _set_cell_padding(lc, top=160, bottom=160, left=220, right=100)
    _set_cell_padding(rc, top=160, bottom=160, left=100, right=220)

    p1 = lc.paragraphs[0]
    p1.paragraph_format.space_before = Pt(0); p1.paragraph_format.space_after = Pt(2)
    r1 = p1.add_run(biz)
    r1.font.name = "Arial"; r1.font.size = Pt(20); r1.font.bold = True
    r1.font.color.rgb = WHITE

    p2 = lc.add_paragraph()
    p2.paragraph_format.space_before = Pt(2); p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(meta["Tagline"])
    r2.font.name = "Arial"; r2.font.size = Pt(9.5)
    r2.font.color.rgb = GRAY

    p_pre = rc.paragraphs[0]
    p_pre.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_pre.paragraph_format.space_before = Pt(0); p_pre.paragraph_format.space_after = Pt(1)
    r_pre = p_pre.add_run("POWERED BY")
    r_pre.font.name = "Arial"; r_pre.font.size = Pt(7); r_pre.font.bold = True
    r_pre.font.color.rgb = WHITE

    p3 = rc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p3.paragraph_format.space_before = Pt(0); p3.paragraph_format.space_after = Pt(2)
    r3 = p3.add_run("EchoFrame")
    r3.font.name = "Arial"; r3.font.size = Pt(26); r3.font.bold = True
    r3.font.color.rgb = AMBER

    p4 = rc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p4.paragraph_format.space_before = Pt(0); p4.paragraph_format.space_after = Pt(0)
    r4 = p4.add_run("Business intelligence, not accounting.")
    r4.font.name = "Arial"; r4.font.size = Pt(8); r4.font.italic = True
    r4.font.color.rgb = WHITE

    sp = doc.add_paragraph()
    sp.paragraph_format.space_before = Pt(4); sp.paragraph_format.space_after = Pt(0)


def _redacted_banner(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(6)
    r1 = p.add_run("REDACTED SAMPLE  ")
    r1.font.name = "Arial"; r1.font.size = Pt(8.5); r1.font.bold = True
    r1.font.color.rgb = AMBER
    r2 = p.add_run("— Fictional business and competitors. Prices, ratings, and promotions are "
                   "illustrative only and do not reflect any real establishment.")
    r2.font.name = "Arial"; r2.font.size = Pt(8.5); r2.font.italic = True
    r2.font.color.rgb = GRAY


def _build_kpis(doc, kpis):
    tbl = doc.add_table(rows=1, cols=3)
    tbl.allow_autofit = False
    _hide_table_borders(tbl)
    col_w = int(7.4 * 1440 / 3)
    for cell, (label, value, sub, val_color) in zip(tbl.rows[0].cells, kpis):
        _set_cell_width_dxa(cell, col_w)
        _set_cell_bg(cell, LTBLUE_HEX)
        _set_cell_padding(cell, top=90, bottom=90, left=220, right=140)
        tcPr = cell._tc.get_or_add_tcPr()
        tcBdr = OxmlElement("w:tcBorders")
        acc = OxmlElement("w:left")
        acc.set(qn("w:val"), "single"); acc.set(qn("w:sz"), "18"); acc.set(qn("w:color"), AMBER_HEX)
        tcBdr.append(acc); tcPr.append(tcBdr)

        pl = cell.paragraphs[0]
        pl.paragraph_format.space_before = Pt(0); pl.paragraph_format.space_after = Pt(4)
        rl = pl.add_run(label.upper())
        rl.font.name = "Arial"; rl.font.size = Pt(10); rl.font.bold = True
        rl.font.color.rgb = GRAY

        pv = cell.add_paragraph()
        pv.paragraph_format.space_before = Pt(0); pv.paragraph_format.space_after = Pt(4)
        rv = pv.add_run(value)
        rv.font.name = "Arial"; rv.font.size = Pt(22); rv.font.bold = True
        rv.font.color.rgb = NAVY

        ps = cell.add_paragraph()
        ps.paragraph_format.space_before = Pt(0); ps.paragraph_format.space_after = Pt(0)
        rs = ps.add_run(sub)
        rs.font.name = "Arial"; rs.font.size = Pt(10); rs.font.bold = True
        rs.font.color.rgb = val_color


def _build_comp_table(doc, rows):
    headers = ["Competitor", "Key Price\n(Lg 1-Topping)", "Rating / Reviews", "Current Promotion", "Change vs\nLast Month"]
    widths = [0.27, 0.13, 0.16, 0.30, 0.14]
    tbl = doc.add_table(rows=1 + len(rows), cols=5)
    tbl.allow_autofit = False
    _hide_table_borders(tbl)
    total = int(7.4 * 1440)

    # Header row
    for cell, htext, w in zip(tbl.rows[0].cells, headers, widths):
        _set_cell_width_dxa(cell, int(total * w))
        _set_cell_bg(cell, NAVY_HEX)
        _set_cell_padding(cell, top=70, bottom=70, left=120, right=100)
        _set_cell_valign(cell, "center")
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
        r = p.add_run(htext)
        r.font.name = "Arial"; r.font.size = Pt(8.5); r.font.bold = True
        r.font.color.rgb = WHITE

    for i, (name, price, rating, promo, change, change_color) in enumerate(rows):
        cells = tbl.rows[i + 1].cells
        for c, w in zip(cells, widths):
            _set_cell_width_dxa(c, int(total * w))
            _set_cell_bg(c, SLATE_HEX if i % 2 else "FFFFFF")
            _set_cell_padding(c, top=70, bottom=70, left=120, right=100)
            _set_cell_valign(c, "center")

        is_you = "(You)" in name

        def _put(cell, text, *, bold=False, color=NAVY, size=9):
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            r.font.name = "Arial"; r.font.size = Pt(size); r.font.bold = bold
            r.font.color.rgb = color

        _put(cells[0], name, bold=is_you, color=NAVY)
        _put(cells[1], price, bold=True, color=NAVY)
        _put(cells[2], rating, color=NAVY)
        _put(cells[3], promo, color=GRAY if promo == "None active" else NAVY, size=8.5)
        _put(cells[4], change, bold=change != "—", color=change_color)


def _build_alerts(doc, alerts):
    for marker_color, lead, rest in alerts:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Pt(2)
        rm = p.add_run("●  ")
        rm.font.name = "Arial"; rm.font.size = Pt(10); rm.font.bold = True
        rm.font.color.rgb = marker_color
        rl = p.add_run(lead)
        rl.font.name = "Arial"; rl.font.size = Pt(10); rl.font.bold = True
        rl.font.color.rgb = NAVY
        rr = p.add_run(rest)
        rr.font.name = "Arial"; rr.font.size = Pt(10)
        rr.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def _build_one_thing(doc):
    ph = doc.add_paragraph()
    ph.paragraph_format.space_before = Pt(2); ph.paragraph_format.space_after = Pt(4)
    rh = ph.add_run(ONE_THING_HEAD)
    rh.font.name = "Arial"; rh.font.size = Pt(11); rh.font.bold = True
    rh.font.color.rgb = NAVY

    pb = doc.add_paragraph()
    pb.paragraph_format.space_before = Pt(0); pb.paragraph_format.space_after = Pt(5)
    rb = pb.add_run(ONE_THING_BODY)
    rb.font.name = "Arial"; rb.font.size = Pt(10)

    for i, step in enumerate(ONE_THING_STEPS, 1):
        ps = doc.add_paragraph()
        ps.paragraph_format.space_before = Pt(1); ps.paragraph_format.space_after = Pt(1)
        rn = ps.add_run(f"{i}.  ")
        rn.font.name = "Arial"; rn.font.size = Pt(10); rn.font.bold = True
        rn.font.color.rgb = NAVY
        rs = ps.add_run(step)
        rs.font.name = "Arial"; rs.font.size = Pt(10)

    pi = doc.add_paragraph()
    pi.paragraph_format.space_before = Pt(6); pi.paragraph_format.space_after = Pt(0)
    ri = pi.add_run(ONE_THING_IMPACT)
    ri.font.name = "Arial"; ri.font.size = Pt(10); ri.font.italic = True
    ri.font.color.rgb = GREEN

    pr = doc.add_paragraph()
    pr.paragraph_format.space_before = Pt(4); pr.paragraph_format.space_after = Pt(0)
    rlab = pr.add_run("Risk: ")
    rlab.font.name = "Arial"; rlab.font.size = Pt(9.5); rlab.font.bold = True
    rlab.font.color.rgb = AMBER
    rrk = pr.add_run(ONE_THING_RISK)
    rrk.font.name = "Arial"; rrk.font.size = Pt(9.5); rrk.font.italic = True
    rrk.font.color.rgb = GRAY


def build():
    from docx.enum.text import WD_LINE_SPACING
    doc = Document()
    _set_doc_defaults(doc)
    # Tighten for a dense single page
    nf = doc.styles["Normal"].paragraph_format
    nf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    nf.line_spacing = 1.08
    nf.space_after = Pt(4)
    sec = doc.sections[0]
    sec.page_width = Inches(8.5); sec.page_height = Inches(11)
    sec.left_margin = sec.right_margin = Inches(0.55)
    sec.top_margin = sec.bottom_margin = Inches(0.5)

    _build_footer(doc, META["Month"])
    # one-page sample: override cached NUMPAGES "4" -> "1"
    for p in doc.sections[0].footer.paragraphs:
        for t in p._p.iter(qn("w:t")):
            if t.text == "4":
                t.text = "1"

    _build_rival_header(doc, META)
    _redacted_banner(doc)
    _build_kpis(doc, KPIS)

    _dark_section_bar(doc, "COMPETITOR COMPARISON  ·  Large 1-Topping Pizza")
    _build_comp_table(doc, COMP_ROWS)

    _dark_section_bar(doc, "WHAT MOVED THIS WEEK")
    _build_alerts(doc, ALERTS)

    _dark_section_bar(doc, "THE ONE THING TO DO THIS WEEK")
    _build_one_thing(doc)

    out_dir = Path(E.BASE_DIR) / "demo_output"
    out_dir.mkdir(exist_ok=True)
    out_path = out_dir / "RivalScan_Sample_TonysBrickOvenPizza.docx"
    doc.save(out_path)
    print(f"Saved: {out_path}")
    return out_path


if __name__ == "__main__":
    build()
